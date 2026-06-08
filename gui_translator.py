#!/usr/bin/env python3
"""PDF 번역기 — 포맷 보존 · 진행률 표시 · PDF 출력"""

import sys, os, json, threading, time, tempfile
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

if sys.platform == "win32":
    import ctypes
    ctypes.windll.kernel32.SetConsoleOutputCP(65001)
    ctypes.windll.kernel32.SetConsoleCP(65001)
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    HAS_DND = True
except ImportError:
    HAS_DND = False

import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# ── 설정 ──────────────────────────────────────────────────────────────────────
CONFIG_FILE = Path.home() / ".pdf_translator_config.json"

LANGUAGES = [
    ("ko","한국어"), ("en","영어"), ("ja","일본어"),
    ("zh-CN","중국어(간체)"), ("zh-TW","중국어(번체)"),
    ("de","독일어"), ("fr","프랑스어"), ("es","스페인어"),
    ("it","이탈리아어"), ("pt","포르투갈어"), ("ru","러시아어"),
    ("vi","베트남어"), ("th","태국어"), ("id","인도네시아어"),
]

# ── 다크 팔레트 ──────────────────────────────────────────────────────────────
_BG      = "#0f172a"
_SURFACE = "#1e293b"
_BORDER  = "#2d3f55"
_ACCENT  = "#3b82f6"
_ACCENT2 = "#2563eb"
_FG      = "#f1f5f9"
_SUBTLE  = "#64748b"
_GREEN   = "#22c55e"
_RED     = "#ef4444"
_YELLOW  = "#f59e0b"
_INFO    = "#60a5fa"
_LOG_BG  = "#080f1c"
_DROP_BG = "#111f35"
_DROP_HL = "#182d4a"

# ─────────────────────────────────────────────────────────────────────────────
LINE_HEIGHT_RATIO       = 1.75  # 한국어 출판 표준 행간
LINE_HEIGHT_RATIO_LATIN = 1.65  # 라틴 계열 언어(영어·독어 등) 행간
PARA_GAP_RATIO          = 0.6   # 단락 사이 추가 여백 (행간의 배수)

KO_FONT_CANDIDATES = [
    "C:/Windows/Fonts/KoPubBatangMedium.ttf",
    "C:/Windows/Fonts/NanumMyeongjo.ttf",
    "C:/Windows/Fonts/HANBatang.ttf",
    "C:/Windows/Fonts/batang.ttc",
    "C:/Windows/Fonts/malgun.ttf",
]

_CJK_LANG_PREFIXES = {"ko", "ja", "zh"}


def _is_cjk_lang(lang_code: str) -> bool:
    return lang_code.lower().split("-")[0] in _CJK_LANG_PREFIXES


def _select_output_font(tgt_lang: str):
    """대상 언어에 맞는 fitz.Font 반환. CJK → 한국어 TTF, 그 외 → Helvetica."""
    import fitz as _fitz
    if _is_cjk_lang(tgt_lang):
        path = next((p for p in KO_FONT_CANDIDATES if os.path.exists(p)), None)
        return _fitz.Font(fontfile=path) if path else _fitz.Font("cjk")
    return _fitz.Font("helv")


_tl_cache: dict = {}  # (text, fontsize) → float  단어/문자 폭 캐시


def _wrap_lines(text, bbox_w, font, fontsize):
    """텍스트를 bbox_w 폭으로 줄바꿈. URL 등 긴 단어는 문자 단위 분할."""
    cache = _tl_cache

    def _tl(t):
        k = (font.name, t, fontsize)
        v = cache.get(k)
        if v is None:
            cache[k] = v = font.text_length(t, fontsize)
        return v

    space_w = _tl(" ")

    def _char_split(w):
        parts, part, pw = [], "", 0.0
        for ch in w:
            cw = _tl(ch)
            if pw + cw <= bbox_w:
                part += ch; pw += cw
            else:
                if part: parts.append(part)
                part, pw = ch, cw
        if part: parts.append(part)
        return parts

    words = text.split()
    lines, cur_words, cur_w = [], [], 0.0
    for w in words:
        ww = _tl(w)
        extra = (space_w if cur_words else 0.0) + ww
        if cur_w + extra <= bbox_w:
            cur_words.append(w); cur_w += extra
        else:
            if cur_words: lines.append(" ".join(cur_words))
            if ww > bbox_w:
                parts = _char_split(w)
                lines.extend(parts[:-1])
                cur_words = [parts[-1]] if parts else []
                cur_w = _tl(cur_words[0]) if cur_words else 0.0
            else:
                cur_words = [w]; cur_w = ww
    if cur_words:
        lines.append(" ".join(cur_words))
    return lines


def _insert_ko_text(page, bbox, text, font, fontsize, color=(0, 0, 0),
                    lh_ratio=LINE_HEIGHT_RATIO):
    """TextWriter로 bbox 내에 텍스트를 줄바꿈하여 삽입. bbox 초과 시 폰트 축소."""
    import fitz as _fitz
    bw   = max(bbox.x1 - bbox.x0, 1.0)
    bh   = bbox.y1 - bbox.y0
    size = max(fontsize, 6.0)
    lines = _wrap_lines(text, bw, font, size)
    # 1pt 하단 여백을 두어 셀 경계선 겹침 방지
    bh_eff = max(bh - 1.0, bh * 0.9)
    while size > 6.0:
        lines = _wrap_lines(text, bw, font, size)
        bh_eff = max(bh - 1.0, bh * 0.9)
        if len(lines) * size * lh_ratio <= bh_eff or size <= 6.0:
            break
        size -= 0.5
    lh = size * lh_ratio
    y  = bbox.y0 + size
    clip_y = bbox.y1 - size * 0.2  # 디센더가 경계선을 넘지 않도록 클립
    tw = _fitz.TextWriter(page.rect, color=color)
    for line in lines:
        if y > clip_y:
            break
        tw.append((bbox.x0, y), line, font=font, fontsize=size)
        y += lh
    tw.write_text(page)


# ─────────────────────────────────────────────────────────────────────────────
def load_config():
    try:
        if CONFIG_FILE.exists():
            return json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}

def save_config(data):
    CONFIG_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def normalize_drop(raw):
    p = raw.strip()
    if p.startswith("{") and p.endswith("}"):
        p = p[1:-1]
    return p.strip('"').strip("'")

# ── 번역 엔진 ─────────────────────────────────────────────────────────────────
SEP       = "\n◆◆◆\n"   # 번역 후에도 살아남는 구분자
MAX_CHARS = 4800          # Google Translate 단일 요청 한도
WORKERS   = 6             # 병렬 요청 수


def _gt(source: str, target: str):
    from deep_translator import GoogleTranslator
    return GoogleTranslator(source=source, target=target)


def _safe_translate(text: str, source: str, target: str) -> str:
    """단일 텍스트 번역. 실패·None 시 원문 반환."""
    if not text or not text.strip():
        return text
    try:
        r = _gt(source, target).translate(text)
        return r if isinstance(r, str) and r else text
    except Exception:
        return text


def _translate_joined(items: list, source: str, target: str) -> list:
    """
    items: [(global_idx, text), ...]
    여러 단락을 SEP으로 합쳐 1회 API 호출로 번역.
    구분자가 깨지면 절반씩 재귀 분할.
    어떤 상황에서도 len(items) 길이의 리스트 반환.
    """
    if not items:
        return []

    idxs  = [i for i, _ in items]
    texts = [t if isinstance(t, str) else "" for _, t in items]

    # 여러 단락 → SEP 구분 → 1회 요청
    if len(texts) > 1:
        try:
            joined = SEP.join(texts)
            raw    = _gt(source, target).translate(joined)
            if isinstance(raw, str) and raw:
                parts = raw.split(SEP)
                if len(parts) == len(texts):
                    return list(zip(idxs, parts))
        except Exception:
            pass

    # 단일 항목 처리
    if len(items) == 1:
        return [(idxs[0], _safe_translate(texts[0], source, target))]

    # 구분자 불일치 → 절반씩 재귀
    mid = len(items) // 2
    left  = _translate_joined(items[:mid], source, target)
    right = _translate_joined(items[mid:], source, target)
    return left + right


def _build_chunks(non_empty: list) -> list:
    """MAX_CHARS 이하가 되도록 단락 묶음 생성"""
    chunks, cur, cur_len = [], [], 0
    sep_len = len(SEP)
    for item in non_empty:
        text     = item[1] if isinstance(item[1], str) else ""
        text_len = len(text)
        extra    = sep_len if cur else 0
        if cur_len + extra + text_len > MAX_CHARS and cur:
            chunks.append(cur)
            cur, cur_len = [item], text_len
        else:
            cur.append(item)
            cur_len += extra + text_len
    if cur:
        chunks.append(cur)
    return chunks


def translate_paragraphs_parallel(all_texts, source, target, progress_cb, cancel):
    """all_texts: list[str] → list[str]  (순서 보존, 병렬 고속)"""
    # None·비문자열 방어
    safe_texts = [t if isinstance(t, str) else "" for t in all_texts]
    result     = list(safe_texts)
    non_empty  = [(i, t) for i, t in enumerate(safe_texts) if t.strip()]
    if not non_empty:
        return result

    chunks = _build_chunks(non_empty)
    total  = len(non_empty)
    done   = 0
    lock   = threading.Lock()

    def do_chunk(chunk):
        nonlocal done
        if cancel.is_set():
            return
        try:
            pairs = _translate_joined(chunk, source, target)
        except Exception:
            # 어떤 오류도 원문으로 폴백
            pairs = [(i, t) for i, t in chunk]
        if not pairs:
            pairs = [(i, t) for i, t in chunk]
        with lock:
            for idx, text in pairs:
                result[idx] = text if isinstance(text, str) else safe_texts[idx]
            done += len(chunk)
            progress_cb(done, total)

    with ThreadPoolExecutor(max_workers=WORKERS) as ex:
        futures = [ex.submit(do_chunk, c) for c in chunks]
        for f in as_completed(futures):
            if cancel.is_set():
                break
            try:
                f.result()
            except Exception:
                pass  # 개별 청크 실패는 무시 (원문 유지됨)

    return result


# 전자책 폰트 구성
# - 본문(작은 글씨): 바탕  — 세리프, 긴 텍스트 가독성 최적
# - 제목(큰 글씨/굵게): 맑은 고딕 — 깔끔한 산세리프, 시각적 구분
FONT_BODY    = "바탕"
FONT_HEADING = "맑은 고딕"
HEADING_PT   = 16   # 16pt 이상만 제목 폰트


def _set_run_font(run, font_name: str):
    """run의 ascii·hAnsi·eastAsia·cs 폰트를 모두 같은 값으로 설정."""
    from docx.oxml.ns import qn
    run.font.name = font_name
    rPr    = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def _pick_font(para) -> str:
    """
    단락의 폰트를 결정한다.
      - 최대 글씨 크기 >= HEADING_PT(16pt) → 제목 폰트 (맑은 고딕)
      - 나머지 → 본문 폰트 (바탕)
    bold 여부는 판단에서 제외: 본문 내 굵은 글씨도 바탕 유지.
    """
    runs = para.runs
    if not runs:
        return FONT_BODY

    # EMU → pt 변환 (1pt = 12700 EMU)
    sizes_pt = [r.font.size / 12700 for r in runs if r.font.size]
    max_pt   = max(sizes_pt) if sizes_pt else 12.0

    return FONT_HEADING if max_pt >= HEADING_PT else FONT_BODY


MIN_BODY_PT      = 9     # 본문 최소 폰트 크기 (pt)
LINE_SPACING     = 1.25  # 줄 간격 배율
BODY_SPACE_AFTER = 0     # 본문 단락 후 여백 pt
HEAD_SPACE_AFTER = 4     # 제목 단락 후 여백 pt
HEAD_SPACE_BEFORE= 3     # 제목 단락 전 여백 pt

import re as _re
# 각주 마커 패턴: 단어·구두점 바로 뒤에 공백 없이 붙은 1~3자리 숫자
# 예) "사망했습니다.7" → "7" 감지 / "10%" → 미감지
_FOOTNOTE_RE = _re.compile(r'(?<=[^\s\d])(\d{1,3})(?!\d)')


def _split_footnotes(run, font_name: str, font_size_emu: int | None):
    """
    run.text 안에서 각주 숫자를 찾아 위첨자 run으로 분리한다.
    원본 run의 텍스트는 비우고, 단락에 새 run들을 추가해 교체한다.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    para = run._r.getparent()   # <w:p> 요소
    text = run.text
    parts = _FOOTNOTE_RE.split(text)   # [normal, num, normal, num, ...]
    if len(parts) == 1:
        return   # 각주 없음 → 변경 불필요

    # 원본 run의 rPr(서식 정보) 복사 — 폰트·크기 기준으로 사용
    base_rpr = run._r.find(qn("w:rPr"))

    # 원본 run 비우기 (나중에 첫 normal 텍스트를 넣을 것)
    run.text = ""

    def _make_rpr(superscript=False):
        """새 run에 붙일 rPr 생성"""
        rpr = OxmlElement("w:rPr")
        # 폰트 설정
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), font_name)
        rpr.append(rFonts)
        # 크기 설정
        if font_size_emu:
            sz_val = str(int(font_size_emu / 12700 * 2))   # half-pt
            for tag in ("w:sz", "w:szCs"):
                el = OxmlElement(tag)
                el.set(qn("w:val"), sz_val)
                rpr.append(el)
        # 위첨자
        if superscript:
            va = OxmlElement("w:vertAlign")
            va.set(qn("w:val"), "superscript")
            rpr.append(va)
        return rpr

    # 원본 run 바로 뒤에 새 run들을 삽입
    insert_after = run._r

    def _insert_run(txt, is_sup):
        r_el = OxmlElement("w:r")
        r_el.append(_make_rpr(superscript=is_sup))
        t_el = OxmlElement("w:t")
        t_el.text = txt
        if txt.startswith(" ") or txt.endswith(" "):
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_el.append(t_el)
        insert_after.addnext(r_el)
        return r_el

    # parts = [normal₀, num₁, normal₂, num₃, ...]  역순으로 삽입
    for part in reversed(parts):
        if not part:
            continue
        is_sup = bool(_FOOTNOTE_RE.fullmatch(part))
        insert_after = _insert_run(part, is_sup)


def apply_translations(doc, items, translated):
    """
    번역 텍스트를 단락 첫 번째 run에 적용.
    - 각주 숫자를 위첨자(superscript)로 분리
    - 본문: bold 제거 + 크기 1pt 축소
    - 줄 간격: LINE_SPACING 배율 적용
    """
    from docx.enum.text import WD_LINE_SPACING

    for (para, _), new_text in zip(items, translated):
        if not para.text.strip() or not para.runs:
            continue

        first      = para.runs[0]
        first.text = new_text or para.text
        font_name  = _pick_font(para)
        _set_run_font(first, font_name)

        # 전체 run 중 최솟값을 기준 크기로 사용 (첫 run만 큰 경우 방지)
        all_sizes = [r.font.size for r in para.runs if r.font.size]
        ref_size  = min(all_sizes) if all_sizes else first.font.size

        if font_name == FONT_BODY:
            first.bold = False
            if ref_size:
                size_pt     = ref_size / 12700
                new_size_pt = max(MIN_BODY_PT, size_pt - 2)
                first.font.size = int(new_size_pt * 12700)
        else:  # FONT_HEADING: 크기는 최솟값으로 정규화 (bold 유지)
            if ref_size:
                first.font.size = ref_size

        # 나머지 run 비우기
        for run in para.runs[1:]:
            run.text = ""

        # 각주 숫자를 위첨자 run으로 분리
        _split_footnotes(first, font_name, first.font.size)

        # 줄 간격 · 단락 여백 설정 (ref_design 기준)
        from docx.shared import Pt as _Pt
        fmt = para.paragraph_format
        fmt.line_spacing      = LINE_SPACING
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        if font_name == FONT_BODY:
            fmt.space_after  = _Pt(BODY_SPACE_AFTER)
            fmt.space_before = None
        else:
            fmt.space_before = _Pt(HEAD_SPACE_BEFORE)
            fmt.space_after  = _Pt(HEAD_SPACE_AFTER)


def _word_to_pdf(docx_abs: str, pdf_abs: str) -> bool:
    """win32com으로 Word를 호출해 PDF 저장. 성공하면 True."""
    import pythoncom
    import win32com.client as wc

    pythoncom.CoInitialize()
    word = None
    doc  = None
    try:
        word = wc.Dispatch("Word.Application")
        word.Visible        = False
        word.DisplayAlerts  = 0   # 모든 팝업·대화상자 억제 (wdAlertsNone)
        word.AutomationSecurity = 3   # 매크로 비활성화

        doc = word.Documents.Open(
            docx_abs,
            False,   # ConfirmConversions
            True,    # ReadOnly
            False,   # AddToRecentFiles
        )
        # ExportAsFixedFormat : Word 2007+ PDF 전용 메서드 (SaveAs보다 안정적)
        doc.ExportAsFixedFormat(
            OutputFileName=pdf_abs,
            ExportFormat=17,          # wdExportFormatPDF
            OpenAfterExport=False,
            OptimizeFor=0,            # wdExportOptimizeForPrint
            Range=0,                  # wdExportAllDocument
            Item=0,                   # wdExportDocumentContent
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=0,        # wdExportCreateNoBookmarks
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
        return True
    finally:
        if doc is not None:
            try: doc.Close(0)          # 0 = wdDoNotSaveChanges
            except Exception: pass
        if word is not None:
            try: word.Quit(0)
            except Exception: pass
        try: pythoncom.CoUninitialize()
        except Exception: pass


def _docx_to_pdf(docx_path: str, pdf_path: str, log):
    """DOCX → PDF. Word(타임아웃 5분) → LibreOffice → DOCX 폴백 순서로 시도."""
    docx_abs = os.path.abspath(docx_path)
    pdf_abs  = os.path.abspath(pdf_path)

    # ── 방법 1: Microsoft Word (별도 스레드 + 5분 타임아웃) ──────────────────
    log("  Word로 PDF 변환 중...  (문서 크기에 따라 수 초~수 분 소요될 수 있습니다)", "warn")
    success  = threading.Event()
    err_box  = [None]

    def _run_word():
        try:
            if _word_to_pdf(docx_abs, pdf_abs):
                success.set()
        except Exception as e:
            err_box[0] = e

    t = threading.Thread(target=_run_word, daemon=True)
    t.start()
    t.join(timeout=300)   # 5분 대기

    if success.is_set():
        log("  PDF 저장 완료", "ok")
        return

    if t.is_alive():
        log("  Word가 응답하지 않아 강제 종료합니다...", "warn")
        # Word 프로세스 강제 종료
        try:
            import subprocess
            subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"],
                           capture_output=True)
        except Exception:
            pass
    else:
        log(f"  Word 실패: {err_box[0]}", "warn")

    # ── 방법 2: LibreOffice ──────────────────────────────────────────────────
    import subprocess as _sp, shutil as _sh
    lo_candidates = [
        _sh.which("soffice"), _sh.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    lo = next((p for p in lo_candidates if p and os.path.exists(p)), None)
    if lo:
        log("  LibreOffice로 PDF 변환 중...")
        try:
            out_dir = os.path.dirname(pdf_abs)
            _sp.run(
                [lo, "--headless", "--convert-to", "pdf",
                 "--outdir", out_dir, docx_abs],
                check=True, timeout=180,
                stdout=_sp.DEVNULL, stderr=_sp.DEVNULL,
            )
            lo_out = os.path.join(
                out_dir,
                os.path.splitext(os.path.basename(docx_abs))[0] + ".pdf"
            )
            if lo_out != pdf_abs and os.path.exists(lo_out):
                os.replace(lo_out, pdf_abs)
            log("  PDF 저장 완료 (LibreOffice)", "ok")
            return
        except Exception as e2:
            log(f"  LibreOffice 실패: {e2}", "warn")

    # ── 최종 폴백: DOCX 저장 ────────────────────────────────────────────────
    import shutil
    docx_out = str(Path(pdf_abs).with_suffix(".docx"))
    shutil.copy(docx_abs, docx_out)
    log("  Word·LibreOffice 모두 실패 → DOCX로 저장했습니다.", "warn")
    log(f"  파일: {docx_out}", "warn")


def collect_paragraphs(doc):
    items = []
    def add(paras, tag):
        for p in paras:
            items.append((p, tag))
    add(doc.paragraphs, "body")
    for i, tbl in enumerate(doc.tables):
        for r, row in enumerate(tbl.rows):
            for c, cell in enumerate(row.cells):
                add(cell.paragraphs, f"t{i}[{r},{c}]")
    for sec in doc.sections:
        for hdr in (sec.header, sec.footer):
            if hdr:
                add(hdr.paragraphs, "hf")
    return items

# ── 앱 ───────────────────────────────────────────────────────────────────────
class App:
    def __init__(self):
        if HAS_DND:
            class _Root(ctk.CTk, TkinterDnD.DnDWrapper):
                def __init__(self_):
                    ctk.CTk.__init__(self_)
                    self_.TkdndVersion = TkinterDnD._require(self_)
            self.root = _Root()
        else:
            self.root = ctk.CTk()

        W, H = 700, 740
        sw = self.root.winfo_screenwidth()
        sh = self.root.winfo_screenheight()
        self.root.geometry(f"{W}x{H}+{(sw - W) // 2}+{(sh - H) // 2}")
        self.root.minsize(600, 640)
        self.root.title("PDF 번역기")
        self.root.configure(fg_color=_BG)

        self.cfg = load_config()
        self._cancel = threading.Event()
        self._build()
        self._load_cfg()

    # ── UI ────────────────────────────────────────────────────────────────────
    def _build(self):
        r = self.root

        # ── 헤더 ──────────────────────────────────────────────────────────────
        hdr = ctk.CTkFrame(r, fg_color=_SURFACE, corner_radius=14,
                           border_color=_BORDER, border_width=1)
        hdr.pack(fill=tk.X, padx=20, pady=(20, 0))

        hrow = ctk.CTkFrame(hdr, fg_color="transparent")
        hrow.pack(fill=tk.X, padx=22, pady=16)

        icon = tk.Label(hrow, text="⬡", bg=_SURFACE, fg=_ACCENT,
                        font=("Segoe UI Symbol", 30))
        icon.pack(side=tk.LEFT, padx=(0, 14))

        tcol = ctk.CTkFrame(hrow, fg_color="transparent")
        tcol.pack(side=tk.LEFT)
        ctk.CTkLabel(tcol, text="PDF 번역기",
                     font=ctk.CTkFont(family="맑은 고딕", size=19, weight="bold"),
                     text_color=_FG).pack(anchor=tk.W)
        ctk.CTkLabel(tcol, text="포맷을 유지하며 PDF를 자동으로 번역합니다",
                     font=ctk.CTkFont(family="맑은 고딕", size=11),
                     text_color=_SUBTLE).pack(anchor=tk.W)

        # ── 파일 선택 ─────────────────────────────────────────────────────────
        fc = ctk.CTkFrame(r, fg_color=_SURFACE, corner_radius=14,
                          border_color=_BORDER, border_width=1)
        fc.pack(fill=tk.X, padx=20, pady=(12, 0))

        ctk.CTkLabel(fc, text="파일 선택",
                     font=ctk.CTkFont(family="맑은 고딕", size=12, weight="bold"),
                     text_color=_FG).pack(anchor=tk.W, padx=22, pady=(14, 8))

        # 드롭 존
        self.df = tk.Frame(fc, bg=_DROP_BG, height=92, cursor="hand2",
                           highlightthickness=2, highlightbackground=_BORDER)
        self.df.pack(fill=tk.X, padx=20, pady=(0, 12))
        self.df.pack_propagate(False)

        self.di = tk.Label(self.df, text="📄", bg=_DROP_BG,
                           font=("Segoe UI Emoji", 22))
        self.di.pack(pady=(10, 2))
        self.dl = tk.Label(self.df, bg=_DROP_BG, fg=_SUBTLE,
                           font=("맑은 고딕", 10), cursor="hand2",
                           text="PDF를 여기에 끌어다 놓거나   클릭하여 선택")
        self.dl.pack()

        for w in (self.df, self.di, self.dl):
            w.bind("<Button-1>", lambda _: self._pick_in())
            w.bind("<Enter>",    lambda _: self._hover(True))
            w.bind("<Leave>",    lambda _: self._hover(False))
        if HAS_DND:
            for w in (self.df, self.di, self.dl):
                w.drop_target_register(DND_FILES)
                w.dnd_bind("<<Drop>>",      self._drop)
                w.dnd_bind("<<DragEnter>>", lambda _: self._hover(True))
                w.dnd_bind("<<DragLeave>>", lambda _: self._hover(False))

        self.iv = tk.StringVar()
        self.ov = tk.StringVar()
        self._frow(fc, "입력", self.iv, self._pick_in)
        self._frow(fc, "출력", self.ov, self._pick_out)
        tk.Frame(fc, bg=_SURFACE, height=6).pack()

        # ── 번역 언어 ─────────────────────────────────────────────────────────
        lc = ctk.CTkFrame(r, fg_color=_SURFACE, corner_radius=14,
                          border_color=_BORDER, border_width=1)
        lc.pack(fill=tk.X, padx=20, pady=(12, 0))

        ctk.CTkLabel(lc, text="번역 언어",
                     font=ctk.CTkFont(family="맑은 고딕", size=12, weight="bold"),
                     text_color=_FG).pack(anchor=tk.W, padx=22, pady=(14, 8))

        lrow = ctk.CTkFrame(lc, fg_color="transparent")
        lrow.pack(padx=22, pady=(0, 16))

        vals     = [f"{c}  {n}" for c, n in LANGUAGES]
        src_vals = ["auto  자동 감지"] + vals
        combo_kw = dict(
            fg_color=_BG, border_color=_BORDER, width=190,
            button_color=_BORDER, button_hover_color=_ACCENT,
            dropdown_fg_color=_SURFACE, dropdown_text_color=_FG,
            dropdown_hover_color=_BORDER,
            text_color=_FG,
            font=ctk.CTkFont(family="맑은 고딕", size=11),
        )

        self.sc = ctk.CTkComboBox(lrow, values=src_vals, **combo_kw)
        self.sc.set(src_vals[0])
        self.sc.pack(side=tk.LEFT)

        ctk.CTkLabel(lrow, text=" → ",
                     font=ctk.CTkFont(family="맑은 고딕", size=15, weight="bold"),
                     text_color=_ACCENT).pack(side=tk.LEFT, padx=14)

        self.tc = ctk.CTkComboBox(lrow, values=vals,
                                  command=self._tgt_changed, **combo_kw)
        self.tc.set(vals[0])
        self.tc.pack(side=tk.LEFT)

        # ── 실행 버튼 ─────────────────────────────────────────────────────────
        br = ctk.CTkFrame(r, fg_color="transparent")
        br.pack(fill=tk.X, padx=20, pady=(14, 0))

        self.btn = ctk.CTkButton(
            br, text="번역 시작",
            fg_color=_ACCENT, hover_color=_ACCENT2, text_color="#ffffff",
            font=ctk.CTkFont(family="맑은 고딕", size=13, weight="bold"),
            height=46, corner_radius=10, command=self._start)
        self.btn.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))

        self.cbtn = ctk.CTkButton(
            br, text="취소",
            fg_color=_SURFACE, hover_color=_BORDER,
            text_color=_SUBTLE, border_color=_BORDER, border_width=1,
            font=ctk.CTkFont(family="맑은 고딕", size=12),
            height=46, corner_radius=10, width=84, state="disabled",
            command=self._do_cancel)
        self.cbtn.pack(side=tk.LEFT)

        # ── 진행 상황 ─────────────────────────────────────────────────────────
        pc = ctk.CTkFrame(r, fg_color=_SURFACE, corner_radius=14,
                          border_color=_BORDER, border_width=1)
        pc.pack(fill=tk.BOTH, expand=True, padx=20, pady=(14, 20))

        ptop = ctk.CTkFrame(pc, fg_color="transparent")
        ptop.pack(fill=tk.X, padx=22, pady=(14, 6))
        ctk.CTkLabel(ptop, text="진행 상황",
                     font=ctk.CTkFont(family="맑은 고딕", size=12, weight="bold"),
                     text_color=_FG).pack(side=tk.LEFT)
        self.pct_lbl = ctk.CTkLabel(
            ptop, text="0%",
            font=ctk.CTkFont(family="맑은 고딕", size=12, weight="bold"),
            text_color=_ACCENT)
        self.pct_lbl.pack(side=tk.RIGHT)

        self.pbar = ctk.CTkProgressBar(
            pc, fg_color=_BG, progress_color=_ACCENT,
            corner_radius=4, height=7)
        self.pbar.set(0)
        self.pbar.pack(fill=tk.X, padx=22, pady=(0, 6))

        self.status = ctk.CTkLabel(
            pc, text="대기 중",
            font=ctk.CTkFont(family="맑은 고딕", size=10),
            text_color=_SUBTLE)
        self.status.pack(anchor=tk.W, padx=22, pady=(0, 10))

        tk.Frame(pc, bg=_BORDER, height=1).pack(fill=tk.X)

        self.log = ctk.CTkTextbox(
            pc, fg_color=_LOG_BG, text_color=_FG,
            font=ctk.CTkFont(family="맑은 고딕", size=12),
            corner_radius=0, border_width=0, wrap="word")
        self.log.pack(fill=tk.BOTH, expand=True)

        tw = self.log._textbox
        tw.tag_configure("ok",   foreground=_GREEN)
        tw.tag_configure("err",  foreground=_RED)
        tw.tag_configure("info", foreground=_INFO)
        tw.tag_configure("warn", foreground=_YELLOW)
        tw.tag_configure("ts",   foreground=_SUBTLE)
        self.log.configure(state="disabled")

    def _frow(self, parent, label, var, cmd):
        row = ctk.CTkFrame(parent, fg_color="transparent")
        row.pack(fill=tk.X, padx=20, pady=(0, 6))
        ctk.CTkLabel(row, text=label, width=40, anchor=tk.W,
                     font=ctk.CTkFont(family="맑은 고딕", size=11),
                     text_color=_SUBTLE).pack(side=tk.LEFT)
        ctk.CTkEntry(row, textvariable=var,
                     fg_color=_BG, border_color=_BORDER, text_color=_FG,
                     font=ctk.CTkFont(family="맑은 고딕", size=10)
                     ).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(8, 8))
        ctk.CTkButton(row, text="찾기", width=58,
                      fg_color=_SURFACE, hover_color=_BORDER,
                      text_color=_FG, border_color=_BORDER, border_width=1,
                      font=ctk.CTkFont(family="맑은 고딕", size=11),
                      height=30, corner_radius=6,
                      command=cmd).pack(side=tk.LEFT)

    # ── 헬퍼 ─────────────────────────────────────────────────────────────────
    def _load_cfg(self):
        if "target_lang" in self.cfg:
            code = self.cfg["target_lang"]
            for c, n in LANGUAGES:
                if c == code:
                    self.tc.set(f"{c}  {n}"); break
        if "source_lang" in self.cfg:
            v = self.cfg["source_lang"]
            if v == "auto":
                self.sc.set("auto  자동 감지")
            else:
                for c, n in LANGUAGES:
                    if c == v:
                        self.sc.set(f"{c}  {n}"); break

    def _src(self):
        v = self.sc.get()
        return "auto" if v.startswith("auto") else v.split()[0]

    def _tgt(self):
        return self.tc.get().split()[0]

    def _hover(self, on):
        bg = _DROP_HL if on else _DROP_BG
        bd = _ACCENT  if on else _BORDER
        for w in (self.df, self.di, self.dl):
            w.config(bg=bg)
        self.df.config(highlightbackground=bd)

    def _set_in(self, path):
        self.iv.set(path)
        p = Path(path)
        self.ov.set(str(p.parent / f"{p.stem}_{self._tgt()}.pdf"))
        self.dl.config(text=p.name, fg=_FG)
        self.di.config(text="✅")
        self.df.config(highlightbackground=_ACCENT)

    def _drop(self, e):
        p = normalize_drop(e.data)
        if not p.lower().endswith(".pdf"):
            self._log("PDF 파일만 지원합니다.", "warn"); return
        self._set_in(p)

    def _tgt_changed(self, _=None):
        if self.iv.get():
            p = Path(self.iv.get())
            self.ov.set(str(p.parent / f"{p.stem}_{self._tgt()}.pdf"))

    def _pick_in(self):
        p = filedialog.askopenfilename(title="번역할 PDF 선택",
                                       filetypes=[("PDF", "*.pdf"), ("모두", "*.*")])
        if p: self._set_in(p)

    def _pick_out(self):
        p = filedialog.asksaveasfilename(title="출력 파일", defaultextension=".pdf",
                                         filetypes=[("PDF", "*.pdf")])
        if p: self.ov.set(p)

    def _log(self, msg, tag=""):
        ts = time.strftime("[%H:%M:%S] ")
        def _w():
            tw = self.log._textbox
            tw.configure(state="normal")
            tw.insert(tk.END, ts, "ts")
            tw.insert(tk.END, msg + "\n", tag if tag else ())
            tw.see(tk.END)
            tw.configure(state="disabled")
        self.root.after(0, _w)

    # ── 진행률 애니메이션 (main 스레드 루프) ─────────────────────────────────
    def _set_pct(self, pct, status=""):
        """백그라운드 스레드에서 호출. 타겟값만 갱신 (root.after 없음)."""
        self._pct_target = float(pct)
        if status:
            self._pct_status = status

    def _start_anim(self):
        """번역 시작 시 main 스레드에서 1회 호출."""
        self._pct_target = 0.0
        self._pct_status = "시작 중..."
        self._anim_on    = True
        self._anim_tick()

    def _stop_anim(self):
        self._anim_on = False
        target = getattr(self, '_pct_target', 0.0)
        self.pbar.set(target / 100)
        self.pct_lbl.configure(text=f"{int(round(target))}%")

    def _anim_tick(self):
        if not getattr(self, '_anim_on', False):
            return
        target  = getattr(self, '_pct_target', 0.0)
        status  = getattr(self, '_pct_status',  "")
        current = self.pbar.get() * 100

        if status:
            self.status.configure(text=status)
            self._pct_status = ""

        diff = target - current
        if abs(diff) > 0.3:
            nv = current + diff * 0.15
            self.pbar.set(nv / 100)
            self.pct_lbl.configure(text=f"{int(round(nv))}%")

        self.root.after(16, self._anim_tick)  # 60fps 루프

    def _set_running(self, on):
        self.btn.configure(state="disabled" if on else "normal")
        self.cbtn.configure(
            state="normal" if on else "disabled",
            text_color=_FG if on else _SUBTLE)
        if not on:
            self._stop_anim()

    def _do_cancel(self):
        self._cancel.set()
        self._log("취소 요청...", "warn")

    # ── 번역 실행 ─────────────────────────────────────────────────────────────
    def _start(self):
        inp = self.iv.get().strip()
        out = self.ov.get().strip()
        if not inp:
            messagebox.showerror("오류", "PDF 파일을 선택하세요."); return
        if not Path(inp).exists():
            messagebox.showerror("오류", f"파일 없음:\n{inp}"); return
        if not out:
            p = Path(inp)
            out = str(p.parent / f"{p.stem}_{self._tgt()}.pdf")
            self.ov.set(out)

        save_config({"target_lang": self._tgt(), "source_lang": self._src()})
        self._cancel.clear()
        self._set_running(True)
        tw = self.log._textbox
        tw.configure(state="normal"); tw.delete("1.0", tk.END); tw.configure(state="disabled")
        self._start_anim()

        threading.Thread(target=self._worker,
                         args=(inp, out, self._src(), self._tgt()),
                         daemon=True).start()

    def _worker(self, inp, out, src, tgt):
        self._log("라이브러리 로딩 중...", "info")
        try:
            import fitz
            from deep_translator import GoogleTranslator  # noqa
        except ImportError as e:
            self._log(f"패키지 오류: {e}", "err")
            self.root.after(0, lambda: self._set_running(False)); return
        self._log("로딩 완료.", "ok")

        try:
            # ── 1단계: 텍스트 추출 ────────────────────────────────────────
            self._set_pct(2, "1/3  텍스트 추출 중...")
            self._log(f"[1/3] {Path(inp).name} — 텍스트 추출", "info")

            doc = fitz.open(inp)
            total_pages = len(doc)
            self._log(f"  총 {total_pages}페이지")

            # 대상 언어에 맞는 폰트 선택 (CJK → 한국어 명조, 그 외 → Helvetica)
            ko_font = _select_output_font(tgt)
            if _is_cjk_lang(tgt):
                _loaded = next((p for p in KO_FONT_CANDIDATES if os.path.exists(p)), None)
                self._log(f"  폰트: {os.path.basename(_loaded) if _loaded else 'CJK 내장'}", "info")
            else:
                self._log("  폰트: Helvetica (라틴 출력)", "info")
            _lh_ratio = LINE_HEIGHT_RATIO if _is_cjk_lang(tgt) else LINE_HEIGHT_RATIO_LATIN

            # 페이지별 block 수집 (block 단위 번역 → 문장 잘림 방지)
            page_blocks = []
            for pidx in range(total_pages):
                page   = doc[pidx]
                raw_blks = page.get_text("dict",
                                         flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                result = []
                for blk in raw_blks:
                    if blk.get("type") != 0:
                        continue
                    spans = []
                    for line in blk["lines"]:
                        for sp in line["spans"]:
                            clean = " ".join(sp["text"].split())
                            if clean:
                                spans.append({**sp, "text": clean})
                    if not spans:
                        continue
                    result.append({
                        "text":  " ".join(sp["text"] for sp in spans),
                        "spans": spans,
                        "bbox":  fitz.Rect(blk["bbox"]),
                    })
                page_blocks.append(result)

            all_texts = [b["text"] for pb in page_blocks for b in pb]
            self._log(f"  블록 {len(all_texts)}개 발견")
            self._set_pct(10, "2/3  번역 중...")

            if self._cancel.is_set():
                raise InterruptedError()

            # ── 2단계: 번역 ────────────────────────────────────────────────
            self._log(f"[2/3] 번역 중  ({src} → {tgt})...", "info")

            def prog_cb(done, total):
                self._set_pct(10 + done / max(total, 1) * 75,
                              f"2/3  번역 중... {done}/{total}")

            translated = translate_paragraphs_parallel(
                all_texts, src, tgt, prog_cb, self._cancel)

            if self._cancel.is_set():
                raise InterruptedError()

            self._set_pct(85, "3/3  텍스트 삽입 중...")

            # ── 3단계: 텍스트 교체 & 저장 ──────────────────────────────────
            self._log("[3/3] 번역 텍스트 삽입 중...", "info")

            # ── 사전 수집: 이미지 블록 위치 & 특수 페이지 판별 ──────────────
            SPECIAL_CHARS = 400   # 원본 텍스트 이 미만 → 표지·챕터 오프너 등 특수 페이지
            page_img_rects  = []  # 각 페이지의 이미지 bbox 목록
            page_draw_zones = []  # 각 페이지의 벡터 드로잉 union bbox
            special_pages   = set()

            for pidx in range(total_pages):
                self._set_pct(85 + pidx / total_pages * 5,
                              f"3/3  분석 중... {pidx+1}/{total_pages}p")
                raw_blks = doc[pidx].get_text(
                    "dict",
                    flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_IMAGES
                )["blocks"]
                imgs = [fitz.Rect(b["bbox"]) for b in raw_blks if b.get("type") == 1]
                page_img_rects.append(imgs)

                # 벡터 드로잉 영역 수집 (500pt² 이상, 배경 fill 80% 미만, 얇은 가로줄 제외)
                drawings = doc[pidx].get_drawings()
                _page_rect = doc[pidx].rect
                _page_area = _page_rect.get_area()
                sig = []
                for _d in drawings:
                    _r = fitz.Rect(_d["rect"])
                    # 얇은 가로줄(표 행 배경): 높이 < 30pt 이고 폭이 높이의 5배 초과 → 제외
                    if _r.height < 30 and _r.width > _r.height * 5:
                        continue
                    _clipped = (_r & _page_rect).get_area()
                    if 500 <= _clipped < _page_area * 0.8:
                        sig.append(_r)
                if sig:
                    dz = sig[0]
                    for r in sig[1:]: dz |= r
                    page_draw_zones.append(dz)
                else:
                    page_draw_zones.append(fitz.Rect())

                total_chars = sum(len(blk["text"]) for blk in page_blocks[pidx])
                if total_chars < SPECIAL_CHARS:
                    special_pages.add(pidx)

            # 배경 이미지 판별: 텍스트 블록 3개 이상 겹치면 배경(표 행 음영 등) → inplace·skip 제외
            _BG_THRESH = 3
            background_imgs = []
            for pidx in range(total_pages):
                bg = set()
                for i_idx, img_r in enumerate(page_img_rects[pidx]):
                    cnt = sum(
                        1 for blk in page_blocks[pidx]
                        if (blk["bbox"] & img_r).get_area() > blk["bbox"].get_area() * 0.4
                    )
                    if cnt >= _BG_THRESH:
                        bg.add(i_idx)
                background_imgs.append(bg)

            # 원본 텍스트 일괄 제거
            for pidx in range(total_pages):
                self._set_pct(90 + pidx / total_pages * 7,
                              f"3/3  원문 제거 중... {pidx+1}/{total_pages}p")
                page  = doc[pidx]
                pblks = page_blocks[pidx]
                if not pblks:
                    continue
                for blk in pblks:
                    for sp in blk["spans"]:
                        page.add_redact_annot(fitz.Rect(sp["bbox"]), fill=None)
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

            # 배경 이미지 white cover: 이미지에 구워진 원본 텍스트 픽셀 제거
            for pidx in range(total_pages):
                bg_list = [page_img_rects[pidx][i] for i in background_imgs[pidx]]
                if bg_list:
                    page = doc[pidx]
                    for img_r in bg_list:
                        page.add_redact_annot(img_r, fill=(1, 1, 1))
                    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)

            # ── 문서 전체 컬럼 경계 도출 ──────────────────────────────────────
            # 페이지마다 min(x0)를 쓰면 헤더/푸터(x≈18) 블록이 컬럼 기준을 오염.
            # 전체 블록의 x0 최빈값을 문서 단위로 한 번만 계산해 사용.
            from collections import Counter as _Ctr
            _all_x0 = [round(blk["bbox"].x0) for pb in page_blocks for blk in pb]
            _all_x1 = [round(blk["bbox"].x1) for pb in page_blocks for blk in pb]
            doc_ax0 = _Ctr(_all_x0).most_common(1)[0][0] if _all_x0 else 50
            doc_ax1 = _Ctr(_all_x1).most_common(1)[0][0] if _all_x1 else 550

            def _content_area(pblks):
                if not pblks:
                    return None
                bboxes = [blk["bbox"] for blk in pblks]
                return (doc_ax0, doc_ax1,
                        min(b.y0 for b in bboxes), max(b.y1 for b in bboxes))

            page_areas = [_content_area(page_blocks[p]) for p in range(total_pages)]
            valid_areas = [a for a in page_areas if a]
            def _med(lst): s = sorted(lst); return s[len(s) // 2]
            std_area = (doc_ax0, doc_ax1,
                        _med([a[2] for a in valid_areas]),
                        _med([a[3] for a in valid_areas]),
                        ) if valid_areas else (50, 550, 50, 750)

            # 번역 아이템: 정규 큐 / 특수 페이지 맵 / 인플레이스 맵으로 분리
            regular_q   = []          # 일반 페이지 아이템 (리플로우 대상)
            special_map = {}          # {pidx: [items]} — 특수 페이지 전용
            inplace_map = {}          # {pidx: [(bbox, text, size, color)]} — 드로잉 영역 내 텍스트

            def _is_inplace(blk_bbox, pidx, blk=None):
                """블록이 그래픽(이미지·벡터 드로잉) 영역과 40% 이상 겹치면 인플레이스 처리.
                블록 전체 bbox가 못 잡는 경우 스팬 단위로도 검사."""
                blk_area = blk_bbox.get_area()
                if blk_area <= 0:
                    return False
                for img_r in page_img_rects[pidx]:
                    if (blk_bbox & img_r).get_area() > blk_area * 0.4:
                        return True
                    # 블록이 이미지 경계에 걸쳐있을 때 스팬 단위 검사
                    if blk:
                        for ln in blk.get("lines", []):
                            for sp in ln.get("spans", []):
                                sp_r = fitz.Rect(sp["bbox"])
                                sp_a = sp_r.get_area()
                                if sp_a > 0 and (sp_r & img_r).get_area() > sp_a * 0.5:
                                    return True
                if pidx < len(page_draw_zones):
                    dz = page_draw_zones[pidx]
                    if not dz.is_empty and (blk_bbox & dz).get_area() > blk_area * 0.4:
                        return True
                return False

            b_idx = 0
            for pidx in range(total_pages):
                for blk in page_blocks[pidx]:
                    t_text = translated[b_idx]; b_idx += 1
                    if not t_text or not t_text.strip():
                        continue
                    spans = blk["spans"]
                    _sz_w = {}
                    for s in spans:
                        k = round(s["size"], 1)
                        _sz_w[k] = _sz_w.get(k, 0) + len(s["text"])
                    size  = max(_sz_w, key=_sz_w.get)
                    raw_c = spans[0].get("color", 0)
                    color = (((raw_c >> 16) & 0xFF) / 255,
                             ((raw_c >>  8) & 0xFF) / 255,
                             (raw_c & 0xFF) / 255) if isinstance(raw_c, int) \
                             else (raw_c or (0, 0, 0))
                    item = {"text": t_text, "size": size, "color": color}
                    if pidx in special_pages:
                        special_map.setdefault(pidx, []).append(item)
                    elif _is_inplace(blk["bbox"], pidx, blk):
                        inplace_map.setdefault(pidx, []).append(
                            (blk["bbox"], t_text, size, color)
                        )
                    else:
                        regular_q.append(item)

            # 본문 크기 정규화: 최빈 크기를 body_size로 정의하고
            # body_size ±2pt 이내 항목을 통일 (chapter 내 글자 불규칙 방지)
            if regular_q:
                from collections import Counter
                size_counts = Counter(round(item["size"]) for item in regular_q)
                body_size = float(size_counts.most_common(1)[0][0])
                for item in regular_q:
                    if abs(item["size"] - body_size) <= 2.0:
                        item["size"] = body_size

            # 페이지 리플로우
            page_w = doc[0].rect.width
            page_h = doc[0].rect.height

            def _get_area(pidx):
                a = page_areas[pidx] if pidx < len(page_areas) and page_areas[pidx] \
                    else std_area
                ax0, ax1, ay0, ay1 = a
                return ax0, ax1, ay0, ay1

            def _place_items(page, ax0, ax1, ay0, ay1, queue, pidx_=None):
                """queue에서 페이지에 들어갈 만큼 삽입. 잔여는 queue[0]에 남김.
                이미지·draw_zone 위를 지나는 y는 해당 영역 아래로 건너뜀."""
                col = fitz.Rect(ax0, ay0, ax1, ay1)
                skips = sorted(
                    [r for r in (page_img_rects[pidx_] if pidx_ is not None
                                 and pidx_ < len(page_img_rects) else [])
                     if r.intersects(col)],
                    key=lambda r: r.y0
                )
                # draw_zone(표·다이어그램)도 skip 대상에 포함
                if pidx_ is not None and pidx_ < len(page_draw_zones):
                    dz = page_draw_zones[pidx_]
                    if not dz.is_empty and dz.intersects(col):
                        skips.append(dz)
                        skips.sort(key=lambda r: r.y0)

                def _skip(y, sz=0):
                    """텍스트 박스 [y, y+sz]가 이미지 rect와 겹치면 rect 아래로 이동."""
                    changed = True
                    while changed:
                        changed = False
                        for r in skips:
                            if y < r.y1 and y + sz >= r.y0:
                                y = r.y1 + 4
                                changed = True
                    return y

                y    = _skip(ay0)
                tws  = {}
                bw   = max(ax1 - ax0, 1.0)
                while queue:
                    item = queue[0]
                    size = item["size"]
                    lh   = size * _lh_ratio
                    gap  = size * PARA_GAP_RATIO
                    y = _skip(y, lh)
                    if y + size > ay1:
                        break
                    lines    = _wrap_lines(item["text"], bw, ko_font, size)
                    color    = item["color"]
                    if color not in tws:
                        tws[color] = fitz.TextWriter(page.rect, color=color)
                    tw       = tws[color]
                    rendered = 0
                    for line in lines:
                        y = _skip(y, lh)
                        if y + size > ay1:
                            break
                        tw.append((ax0, y + size), line, font=ko_font, fontsize=size)
                        y += lh; rendered += 1
                    if rendered < len(lines):
                        queue[0] = {**item, "text": " ".join(lines[rendered:])}
                        break
                    else:
                        queue.pop(0); y += gap
                for tw in tws.values():
                    tw.write_text(page)

            pidx = 0
            while True:
                if not regular_q and not special_map:
                    break
                if pidx >= len(doc):
                    if not regular_q:
                        break
                    doc.new_page(width=page_w, height=page_h)

                page = doc[pidx]
                ax0, ax1, ay0, ay1 = _get_area(pidx)

                self._set_pct(97 + min(pidx, total_pages) / max(total_pages, 1) * 2,
                              f"3/3  텍스트 삽입 중... {min(pidx+1, total_pages)}/{total_pages}p")
                if pidx in special_pages and pidx < total_pages:
                    # 특수 페이지: 자기 아이템만 배치, overflow 받지 않음
                    page_items = special_map.pop(pidx, [])
                    _place_items(page, ax0, ax1, ay0, ay1, page_items, pidx_=pidx)
                else:
                    # 일반 페이지: 정규 큐에서 리플로우
                    _place_items(page, ax0, ax1, ay0, ay1, regular_q, pidx_=pidx)

                pidx += 1

            # 인플레이스 셀의 이미지 픽셀에서 원문 한글 제거
            # PDF_REDACT_IMAGE_PIXELS: 래스터 이미지 픽셀까지 실제로 덮어씀
            for ip_pidx, items in inplace_map.items():
                page = doc[ip_pidx]
                for bbox, _t, _s, _c in items:
                    page.add_redact_annot(fitz.Rect(bbox), fill=(1, 1, 1))
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)

            # 드로잉 영역 내 텍스트를 원래 위치에 삽입 (인플레이스)
            for ip_pidx, items in inplace_map.items():
                page = doc[ip_pidx]
                for bbox, text, size, color in items:
                    _insert_ko_text(page, bbox, text, ko_font, size, color,
                                    lh_ratio=_lh_ratio)

            self._set_pct(99, "3/3  저장 중...")
            doc.save(out, deflate=True, garbage=4)
            doc.close()

            self._set_pct(100, "완료!")
            self._log(f"\n저장: {out}", "ok")
            self._log("번역 완료", "ok")
            self.root.after(0, lambda: messagebox.showinfo(
                "완료", f"번역 완료!\n\n{out}"))

        except InterruptedError:
            self._log("취소됨.", "warn")
            self._set_pct(0, "취소됨")
        except Exception as e:
            msg = str(e)
            self._log(f"오류: {msg}", "err")
            self.root.after(0, lambda m=msg: messagebox.showerror("오류", m))
        finally:
            self.root.after(0, lambda: self._set_running(False))

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    App().run()
